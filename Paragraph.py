# -- coding: UTF-8 --
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

_MAPPING = (
    u'零', u'一', u'二', u'三', u'四', u'五', u'六', u'七', u'八', u'九', u'十', u'十一', u'十二', u'十三', u'十四', u'十五', u'十六', u'十七',
    u'十八', u'十九')
_P0 = (u'', u'十', u'百', u'千',)
_S4 = 10 ** 4

def num_to_chinese4(num):
    assert (0 <= num and num < _S4)
    if num < 20:
        return _MAPPING[num]
    else:
        lst = []
        while num >= 10:
            lst.append(num % 10)
            num = num / 10
        lst.append(num)
        c = len(lst)  # 位数
        result = u''
        for idx, val in enumerate(lst):
            val = int(val)
            if val != 0:
                result += _P0[idx] + _MAPPING[val]
                if idx < c - 1 and lst[idx + 1] == 0:
                    result += u'零'
        return result[::-1]


def add_paragraph(document, text, style='FangSong_GB2312', align='left'):
    paragraph = document.add_paragraph()
    if align == 'center':
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        paragraph_format = paragraph.paragraph_format
        paragraph_format.first_line_indent = Pt(28)
    run = paragraph.add_run(text, style=style)

    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)


def paragraph(document):
    paragraph = document.paragraphs[0]
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(u'每天30秒电力热讯', style='JianBiaoSong')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(20)

    with open('OriginalFile.txt', 'r', encoding='UTF-8') as file:
        j = 1
        for line in file:
            line = line.replace('✦', '')
            if line.count('#每天30秒电力热讯  '):
                text = '（' + line[12:-1] + '）'
                time = line[12:-1]
                add_paragraph(document, text, align='center')
                add_paragraph(document, '')
            elif line == "【电力快讯】\n":
                text = num_to_chinese4(j) + '、' + '电力快讯'
                add_paragraph(document, text, style='SimHei')
                j = j + 1
                k = 1
                for line in file:
                    line = line.replace('✦', '')
                    if line != '【文章推荐】\n' and line != '\n' and line != '\r':
                        line = line.replace('\n', '')
                        text = '（' + num_to_chinese4(k) + '）' + line
                        add_paragraph(document, text)
                        k = k + 1
                    else:
                        break
            elif line == "【文章推荐】\n":
                text = num_to_chinese4(j) + '、' + '文章推荐'
                add_paragraph(document, text, style='SimHei')
                j = j + 1
                k = 1
                for line in file:
                    line = line.replace('✦', '')
                    if line.count('【综合智慧能源服务展示】'):
                        line = line.replace('\n', '')
                        text1 = '（' + num_to_chinese4(k) + '）' + '综合智慧能源服务展示'
                        add_paragraph(document, text1, 'KaiTi_GB2312')
                        line = line.replace('【综合智慧能源服务展示】', '')
                        index1 = line.find('http')
                        text2 = line[:index1]
                        add_paragraph(document, text2)
                        text3 = '网址：' + line[index1:]
                        add_paragraph(document, text3)
                    elif line.count('【专利开放许可】'):
                        line = line.replace('\n', '')
                        text1 = '（' + num_to_chinese4(k) + '）' + '专利开放许可'
                        add_paragraph(document, text1, 'KaiTi_GB2312')
                        line = line.replace('【专利开放许可】', '')
                        index1 = line.find('http')
                        text2 = line[:index1]
                        add_paragraph(document, text2)
                        text3 = '网址：' + line[index1:]
                        add_paragraph(document, text3)
                    elif line.count('【电厂名单】'):
                        line = line.replace('\n', '')
                        text1 = '（' + num_to_chinese4(k) + '）' + '电厂名单'
                        add_paragraph(document, text1, 'KaiTi_GB2312')
                        line = line.replace('【电厂名单】', '')
                        index1 = line.find('http')
                        text2 = line[:index1]
                        add_paragraph(document, text2)
                        text3 = '网址：' + line[index1:]
                        add_paragraph(document, text3)
                    elif line.count('【文字实录】'):
                        line = line.replace('\n', '')
                        text1 = '（' + num_to_chinese4(k) + '）' + '文字实录'
                        add_paragraph(document, text1, 'KaiTi_GB2312')
                        line = line.replace('【文字实录】', '')
                        index1 = line.find('http')
                        text2 = line[:index1]
                        add_paragraph(document, text2)
                        text3 = '网址：' + line[index1:]
                        add_paragraph(document, text3)
                    k = k + 1
            else:
                print('', end='')
    return time
